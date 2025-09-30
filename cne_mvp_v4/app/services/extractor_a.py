
import os, pandas as pd
from typing import List, Dict, Any
import pdfplumber
from docx import Document
from . import edital_parser
def extract_from_xlsx(path: str) -> List[Dict[str, Any]]:
    df = pd.read_excel(path, dtype=str); return df.to_dict(orient="records")
def extract_from_csv(path: str) -> List[Dict[str, Any]]:
    df = pd.read_csv(path, dtype=str, sep=None, engine="python"); return df.to_dict(orient="records")
def extract_from_docx(path: str) -> List[Dict[str, Any]]:
    doc = Document(path); rows = []
    if doc.tables:
        for tbl in doc.tables:
            headers = [cell.text.strip() for cell in tbl.rows[0].cells]
            for row in tbl.rows[1:]:
                r = {headers[i]: row.cells[i].text.strip() for i in range(min(len(headers), len(row.cells)))}
                rows.append(r)
        return rows
    recs, _ = edital_parser.parse_docx(path); return recs
def extract_from_pdf(path: str) -> List[Dict[str, Any]]:
    rows = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables() or []
            for tbl in tables:
                if len(tbl) > 1:
                    headers = [str(x).strip() for x in tbl[0]]
                    for row in tbl[1:]:
                        r = {headers[i]: (str(row[i]).strip() if i < len(row) else "") for i in range(len(headers))}
                        rows.append(r)
    return rows
def extract(path: str) -> List[Dict[str, Any]]:
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in [".xlsx", ".xls"]: return extract_from_xlsx(path)
        if ext == ".csv": return extract_from_csv(path)
        if ext == ".docx": return extract_from_docx(path)
        if ext == ".pdf": return extract_from_pdf(path)
    except Exception as e:
        print(f"[Extractor A] Falha a processar {path}: {e}")
    return []
